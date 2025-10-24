# -*- coding: utf-8 -*-
"""
Генерация ПНР-документа из .docx-источника:
- парсинг входного DOCX → JSON по схеме (через OpenRouter),
- вычисление FERP-количеств,
- генерация разделов (Задачи/Объём/Порядок/Испытания/Методика) через LLM,
- сборка итогового DOCX по шаблону с аккуратной вставкой абзацев и списков,
- кэширование JSON/текстов в _pnr_cache/<doc_basename_hash>/.

Зависимости: python-docx, docxtpl, openai (клиент OpenRouter), jsonschema (опц.).
Переменные окружения:
  - OPENROUTER_API_KEY (обязательно)
  - OPENROUTER_MODEL   (напр., "qwen/qwen2.5-32b-instruct:free" или др.)
"""

from __future__ import annotations

import json
import os
import re
import time
import hashlib
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
from docxtpl import DocxTemplate
from openai import OpenAI
import google.generativeai as genai

import importlib.resources as pkg_resources  # чтение промптов из пакета

# === Доменные данные (из пакета проекта)
from pnr_extractor.schema import JSON_SCHEMA
from pnr_extractor.ferp_reference import FERP_TASKS, FERP_ORDER, UNIT_BY_CODE, NOTE_BY_CODE


# =============================================================================
# 0) Константы и настройки
# =============================================================================

# Эти два лучше задавать в окружении:
#OPENROUTER_API_KEY = "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"
#OPENROUTER_MODEL = "qwen/qwen3-4b:free"  # можно заменить на другую free-модель
#OPENROUTER_MODEL = "z-ai/glm-4.5-air:free"
OPENROUTER_MODEL = "gemini-2.5-flash"
OPENROUTER_API_KEY = "AIzaSyBoRPAe3T5Ve4FyNSJxei9gu8F_FoTCq4w"

GEMINI_MODEL = "gemini-2.5-flash"
GEMINI_API_KEY = "AIzaSyBoRPAe3T5Ve4FyNSJxei9gu8F_FoTCq4w"

# Пути по умолчанию (можно переопределить при вызове main/process_one_docx_file)
DOCX_PATH_DEFAULT = r"/home/kolya/Documents/Work/СС2/"            # файл .docx или папка с .docx
DOCX_TEMPLATE_PATH_DEFAULT = r"/home/kolya/Documents/Work/template.docx"

# Имена файлов промптов в пакете pnr_extractor
PROMPT_MAIN        = "prompt_pnr_extractor.txt"
PROMPT_METHODOLOGY = "prompt_methodology.txt"
PROMPT_TASKS       = "prompt_tasks.txt"
PROMPT_SCOPE       = "prompt_scope.txt"
PROMPT_PROCEDURE   = "prompt_procedure.txt"
PROMPT_TESTS       = "prompt_tests.txt"

# Ретраи LLM
MAX_LLM_RETRIES = 5
LLM_RETRY_DELAY_SEC = 120  # 3 минуты


# =============================================================================
# 1) Утилиты кэша
# =============================================================================

def _safe_basename(path: str) -> str:
    """Безопасное имя файла + короткий SHA1 хеш от абсолютного пути (устраняем коллизии)."""
    base = os.path.splitext(os.path.basename(path))[0]
    safe = re.sub(r'[^A-Za-z0-9._-]+', '_', base)
    h = hashlib.sha1(os.path.abspath(path).encode("utf-8")).hexdigest()[:8]
    return f"{safe}_{h}"

def _cache_root(docx_path: str) -> str:
    """Каталог кэша для конкретного исходного DOCX."""
    root = os.path.join(os.path.dirname(docx_path), "_pnr_cache", _safe_basename(docx_path))
    os.makedirs(root, exist_ok=True)
    return root

def _cache_path(docx_path: str, filename: str) -> str:
    return os.path.join(_cache_root(docx_path), filename)

def gemini_cache_path(docx_path: str) -> str:
    """Для совместимости: путь, куда пишем структурированный JSON."""
    return _cache_path(docx_path, "gemini.json")

def _methodology_cache_path(docx_path: str) -> str:
    return _cache_path(docx_path, "methodology.txt")

def _section_cache_path(docx_path: str, suffix: str) -> str:
    return _cache_path(docx_path, f"{suffix}.txt")


# =============================================================================
# 2) Чтение исходного DOCX и построение промптов
# =============================================================================

def _load_pkg_text(name: str) -> str:
    """Читает текстовый ресурс из пакета pnr_extractor (UTF-8)."""
    return pkg_resources.files("pnr_extractor").joinpath(name).read_text(encoding="utf-8")

def read_docx_text(path: str) -> str:
    """
    Сбор текстового содержимого DOCX для промпта:
    - абзацы
    - таблицы (строки как |-разделённые ячейки)
    """
    doc = Document(path)
    parts: List[str] = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            parts.append(t)
    for tbl in doc.tables:
        for row in tbl.rows:
            cells = [c.text.strip() for c in row.cells]
            row_text = " | ".join([t for t in cells if t])
            if row_text:
                parts.append(row_text)
    return "\n".join(parts)

def build_prompt(docx_text: str, docx_path: str) -> str:
    """
    Подставляет DOCX_PATH, DOCX_TEXT в шаблон PROMPT_MAIN.
    Надёжно экранирует прочие фигурные скобки.
    """
    template = _load_pkg_text(PROMPT_MAIN)

    marker_path = "__DOCX_PATH__MARKER__"
    marker_text = "__DOCX_TEXT__MARKER__"
    template = template.replace("{DOCX_PATH}", marker_path).replace("{DOCX_TEXT}", marker_text)
    template = template.replace("{", "{{").replace("}", "}}")
    template = template.replace(marker_path, "{DOCX_PATH}").replace(marker_text, "{DOCX_TEXT}")

    try:
        return template.format(DOCX_PATH=docx_path, DOCX_TEXT=docx_text)
    except KeyError as e:
        raise RuntimeError(f"Неэкранированный плейсхолдер в промпте: {e!s}")


# =============================================================================
# 3) Нормализация текста (мягкие переносы, спецсимволы)
# =============================================================================

def normalize_word_breaks(raw: str) -> str:
    """Приводим ответы LLM/Word к предсказуемым \n, чистим NBSP/ZWSP/SHY и лишние переносы."""
    if not raw:
        return ""
    t = raw
    t = re.sub(r'\^l', '\n', t)
    t = t.replace('\x0b', '\n').replace('\u2028', '\n').replace('\u2029', '\n\n')
    t = t.replace('\r\n', '\n').replace('\r', '\n')
    t = t.replace('\u00A0', ' ').replace('\u200B', '').replace('\u00AD', '')
    t = re.sub(r'(\w)[\-–—]\n(\w)', r'\1\2', t)      # слияние переносов слова
    t = re.sub(r'\n{3,}', '\n\n', t)
    return t.strip()


# =============================================================================
# 4) OpenRouter: единый клиент и строгий JSON-ответ по схеме
# =============================================================================

class OpenRouterClient:
    """Совместимая обёртка: та же API, но работает через Gemini.
    - messages: список dict({"role": "system"|"user"|"assistant", "content": str})
    - temperature: пробрасывается в generation_config
    - json_mode: при True запрашиваем строгий JSON через response_mime_type
    Ретраи: до MAX_LLM_RETRIES, пауза LLM_RETRY_DELAY_SEC.
    """

    def __init__(self, api_key: str, model: str) -> None:
        if not api_key:
            raise RuntimeError("API-ключ LLM не задан (ожидался GEMINI_API_KEY).")
        self.model = model
        genai.configure(api_key=api_key)

    def _messages_to_prompt(self, messages: List[Dict[str, str]]) -> str:
        """Простая конкатенация ролей в один промпт для Gemini."""
        parts: List[str] = []
        for m in messages or []:
            role = (m.get("role") or "user").strip().lower()
            content = (m.get("content") or "").strip()
            if not content:
                continue
            if role == "system":
                parts.append(f"System:\n{content}\n")
            elif role == "assistant":
                parts.append(f"Assistant:\n{content}\n")
            else:  # user/прочее
                parts.append(f"User:\n{content}\n")
        # Явная просьба отвечать только контентом без лишних рамок
        parts.append("IMPORTANT: Respond with content only. No Markdown fences.")
        return "\n".join(parts).strip()

    def chat(
        self,
        messages: List[Dict[str, str]],
        *,
        temperature: float = 0.0,
        json_mode: bool = False
    ) -> str:
        prompt = self._messages_to_prompt(messages)

        # Конфиг генерации
        generation_config: Dict[str, Any] = {"temperature": float(temperature)}
        if json_mode:
            # Жёстко просим JSON на уровне mime-типа
            generation_config["response_mime_type"] = "application/json"

        last_err: Optional[Exception] = None
        for attempt in range(1, MAX_LLM_RETRIES + 1):
            try:
                model = genai.GenerativeModel(self.model, generation_config=generation_config)
                resp = model.generate_content(prompt)
                content = (resp.text or "").strip()
                if not content:
                    raise RuntimeError("Пустой ответ модели.")
                # Срезаем возможные ```json ... ```
                if content.startswith("```"):
                    content = content.split("\n", 1)[-1]
                    if "\n```" in content:
                        content = content.rsplit("\n```", 1)[0]
                # Иногда модели префиксуют "json:" — уберём
                low = content.lower()
                if low.startswith("json:"):
                    content = content[5:].lstrip()
                return content
            except Exception as e:
                last_err = e
                if attempt < MAX_LLM_RETRIES:
                    print(f"[RETRY] LLM ошибка/пусто (попытка {attempt}/{MAX_LLM_RETRIES}). "
                          f"Ждём {LLM_RETRY_DELAY_SEC}s…")
                    time.sleep(LLM_RETRY_DELAY_SEC)
                else:
                    break

        assert last_err is not None
        raise last_err


def _strip_md_fences(s: str) -> str:
    """Срезает ```json ... ``` ограды и префикс 'json:' при необходимости."""
    t = s.strip()
    if t.startswith("```"):
        t = t.split("\n", 1)[-1]
        if "\n```" in t:
            t = t.rsplit("\n```", 1)[0]
    if t.lower().startswith("json"):
        t = t[4:].lstrip(":").lstrip()
    return t.strip()

def _validate_schema(obj: Dict[str, Any]) -> None:
    """Пытаемся валидировать по JSON_SCHEMA (если установлен jsonschema)."""
    try:
        import jsonschema  # type: ignore
        jsonschema.validate(instance=obj, schema=JSON_SCHEMA)
    except ImportError:
        return
    except Exception as e:
        raise RuntimeError(f"JSON не соответствует схеме: {e}")

def call_structured_extractor(docx_text: str, docx_path: str) -> Dict[str, Any]:
    """
    Извлекает структурированный JSON строго по JSON_SCHEMA.
    Делаем до 5 попыток пакетом: (soft → hard).
    Если пришло, но JSON невалиден/не по схеме — ждём 3 минуты и повторяем пакет.
    """
    client = OpenRouterClient(OPENROUTER_API_KEY, OPENROUTER_MODEL)

    user_prompt = build_prompt(docx_text, docx_path)
    schema_json_str = json.dumps(JSON_SCHEMA, ensure_ascii=False, indent=2)

    SYSTEM = (
        "You extract structured data. Return a SINGLE valid UTF-8 JSON object. "
        "No prose/markdown/comments."
    )
    SCHEMA = (
        "Output MUST strictly match this JSON Schema (draft-07 semantics):\n"
        f"{schema_json_str}\n\n"
        "If a field is unknown, use schema-conformant defaults (empty list/string/0)."
    )
    STRICT = (
        "\n\nIMPORTANT:\n- Return ONLY a JSON object.\n- No ``` fences.\n- No explanations.\n"
    )

    msgs_soft = [
        {"role": "system", "content": SYSTEM},
        {"role": "system", "content": SCHEMA},
        {"role": "user", "content": user_prompt + STRICT},
    ]
    msgs_hard = [
        {"role": "system", "content": SYSTEM},
        {"role": "system", "content": SCHEMA},
        {"role": "user", "content": user_prompt},
        {"role": "user", "content": "Return ONLY a single JSON object conforming to the schema. No markdown."},
    ]

    def _try_once() -> Dict[str, Any]:
        # soft
        try:
            raw = client.chat(msgs_soft, temperature=0.0, json_mode=False)
            raw = _strip_md_fences(raw)
            data = json.loads(raw)
            _validate_schema(data)
            return data
        except Exception:
            pass
        # hard (json mode)
        raw = client.chat(msgs_hard, temperature=0.0, json_mode=True)
        raw = _strip_md_fences(raw)
        data = json.loads(raw)
        _validate_schema(data)
        return data

    last_err: Optional[Exception] = None
    for attempt in range(1, MAX_LLM_RETRIES + 1):
        try:
            return _try_once()
        except Exception as e:
            last_err = e
            if attempt < MAX_LLM_RETRIES:
                print(f"[RETRY] JSON невалиден/не по схеме (попытка {attempt}/{MAX_LLM_RETRIES}). Ждём {LLM_RETRY_DELAY_SEC}s…")
                time.sleep(LLM_RETRY_DELAY_SEC)
            else:
                break

    assert last_err is not None
    raise RuntimeError(f"Не удалось получить валидный JSON от модели. Последняя ошибка: {last_err}")



# =============================================================================
# 5) Подсчёты FERP и форматирование чисел
# =============================================================================

def _classify_items(equipment: List[Dict[str, Any]]) -> Dict[str, int]:
    """
    Грубая классификация элементов оборудования по ключевым группам для расчётов FERP.
    """
    d = {k: 0 for k in [
        'rj45_points', 'optical_patchcords', 'sfp_modules', 'network_patchcords',
        'network_cable_spools', 'power_cable_segments', 'shielded_cable_segments',
        'cabinet', 'ground_busbar', 'din_panel', 'pdu', 'breaker', 'socket_din',
        'switch', 'router', 'ups', 'computer', 'voip_phone', 'cctv', 'ntp_clock',
        'access_control', 'endpoints', 'telemechanics_channels'
    ]}
    for it in equipment or []:
        et = (it.get('equipment_type') or '').lower()
        name = ((it.get('name') or '') + ' ' + (it.get('model_or_code') or '')).lower()
        q = int(it.get('quantity') or 0)

        if et in d:
            d[et] += q
        if et in {'computer', 'voip_phone', 'cctv', 'ntp-clock', 'access_control'}:
            d['endpoints'] += q

        if re.search(r'(розеточ\w+ модул\w+|keystone|розетк\w+)', name) and 'коннектор' not in name:
            d['rj45_points'] += q
        if ('оптическ' in name) and re.search(r'(шнур|патч-?корд).*(sc|lc)', name):
            d['optical_patchcords'] += q
        if 'sfp' in name:
            d['sfp_modules'] += q
        if re.search(r'коммутационн\w+ шнур', name) and ('cat' in name or 'категор' in name):
            d['network_patchcords'] += q
        if re.search(r'кабель витая пара|utp|ftp|s\/ftp|f\/utp', name):
            d['network_cable_spools'] += q
            if re.search(r'(f\/utp|s\/ftp|sf\/utp|экранир)', name):
                d['shielded_cable_segments'] += q
        if re.search(r'(ввг|пввг|пугв|nym|kg|силов)', name):
            d['power_cable_segments'] += q
        if 'сигнализац' in name or 'телемехан' in name:
            d['telemechanics_channels'] += q
    return d

def _compute_ferp_counts(data: Dict[str, Any]) -> Dict[str, Any]:
    """Итоговые количества для таблиц ФЕРп и генерации разделов."""
    eq = data.get('equipment', []) or []
    c = _classify_items(eq)

    scs_points       = c['rj45_points'] if c['rj45_points'] else (c['endpoints'] if c['endpoints'] else (c['network_patchcords'] // 2))
    optical_lines    = max(c['sfp_modules'] // 2, c['optical_patchcords'] // 2)
    optical_points   = c['optical_patchcords'] if c['optical_patchcords'] else (optical_lines * 2)
    lines_comm_total = scs_points + optical_lines

    power_points              = c['ups'] + (1 if c['cabinet'] else 0)
    ground_resistance_points  = 1 if (c['ground_busbar'] or c['cabinet']) else 0
    ground_continuity_all     = c['cabinet'] + c['din_panel'] + c['pdu'] + c['ups'] + c['switch'] + c['router'] + c['computer']
    ground_continuity_equip   = c['ups'] + c['switch'] + c['router'] + c['computer']
    ground_continuity_racks   = c['cabinet']

    return {
        "01-11-010-01": ground_resistance_points,
        "01-11-011-01_all":   ground_continuity_all,
        "01-11-011-01_equip": ground_continuity_equip,
        "01-11-011-01_racks": ground_continuity_racks,
        "01-11-028-01": c['power_cable_segments'],
        "01-06-022-01": scs_points,
        "02-01-002-01": lines_comm_total,
        "02-01-003-01": c['telemechanics_channels'],
        "02-01-004-01": optical_points,
        "01-11-012-01": power_points,
        "01-11-026-01": power_points,
        "01-11-034-01": c['ups'],
        "01-11-011-02": c['shielded_cable_segments'],
        "01-11-029-01": 1 if ground_resistance_points else 0,
        "01-11-036-01": 1 if eq else 0,
    }

def _fmt_ru_decimal(x: float) -> str:
    """0.03 → '0,03'."""
    s = f"{x:.2f}".rstrip("0").rstrip(".")
    return s.replace(".", ",")


# =============================================================================
# 6) DOCX: стили, поиск таблиц, вставка контента
# =============================================================================

def ensure_paragraph_style(doc: Document, style_name: str, *,
                           font_name="Times New Roman", font_size_pt=12,
                           bold=False, left_indent_pt=0):
    """Создаёт/обновляет стиль абзаца (если нет)."""
    styles = doc.styles
    try:
        style = styles[style_name]
    except KeyError:
        style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
    style.base_style = styles["Normal"]
    font = style.font
    font.name = font_name
    font.size = Pt(font_size_pt)
    font.bold = bold
    pf = style.paragraph_format
    pf.left_indent = Pt(left_indent_pt)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    return style

def find_table_by_headers(doc: Document, required_headers: List[str], min_match: int) -> Optional[Any]:
    """Грубый поиск таблицы по заголовку первой строки (минимум совпадений min_match)."""
    req = set(h.strip().lower() for h in required_headers)
    for table in doc.tables:
        if not table.rows:
            continue
        header_cells = [c.text.strip().lower() for c in table.rows[0].cells]
        if len(req.intersection(header_cells)) >= min_match:
            return table
    return None

def strip_num_prefix(s: str) -> str:
    """'1) Текст' → 'Текст'."""
    return re.sub(r'^\s*\d+\s*[.)]\s*', '', s or '').strip()

def insert_text_with_style(doc: Document, placeholder: str, text: str, style_name: str = "Normal") -> None:
    """
    Простейшая «умная» вставка:
      - бьём на блоки по пустой строке,
      - внутри блока распознаём списки: нумерованные (`1.`, `1)`), маркированные (`-`, `•`),
      - иначе — обычные абзацы.
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    for p in doc.paragraphs:
        if placeholder not in p.text:
            continue

        parent = p._element.getparent()
        idx = parent.index(p._element)
        parent.remove(p._element)

        blocks = [b.strip() for b in re.split(r'\n\s*\n', (text or "").strip()) if b.strip()]
        for block in blocks:
            lines = [l.strip() for l in block.splitlines() if l.strip()]
            is_list = all(re.match(r'^(\d+[\.\)]|[-•])\s+', l) for l in lines)

            for line in lines:
                clean = re.sub(r'^(\d+[\.\)]|[-•])\s+', '', line)
                new_p = doc.add_paragraph(clean, style=style_name)

                if is_list:
                    pPr = new_p._element.get_or_add_pPr()
                    numPr = OxmlElement('w:numPr')
                    ilvl = OxmlElement('w:ilvl'); ilvl.set(qn('w:val'), '0')
                    numId = OxmlElement('w:numId'); numId.set(qn('w:val'), '1')  # fallback-нумерация
                    numPr.append(ilvl); numPr.append(numId)
                    pPr.append(numPr)

                parent.insert(idx, new_p._element)
                idx += 1
        break


# =============================================================================
# 7) Генераторы разделов через LLM
# =============================================================================

def _equipment_block_text(equipment: List[Dict[str, Any]]) -> str:
    lines = []
    for it in equipment or []:
        et     = it.get("equipment_type", "")
        vendor = it.get("vendor", "")
        name   = it.get("name", "")
        model  = it.get("model_or_code", "")
        qty    = it.get("quantity", 0)
        ferp_list = ", ".join(
            f"{f.get('code','')} ({f.get('title','')})"
            for f in (it.get("ferp") or [])
            if f.get("code")
        )
        lines.append(f"- {et}: {vendor} {name} {model} — {qty} шт. FERP: {ferp_list}")
    return "\n".join(lines) if lines else "—"

def _counts_block_text(counts: Dict[str, Any]) -> str:
    lines = [f"{code}: {counts.get(code, 0)}" for code in FERP_ORDER]
    for k in ("01-11-011-01_all", "01-11-011-01_equip", "01-11-011-01_racks"):
        if k in counts:
            lines.append(f"{k}: {counts[k]}")
    return "\n".join(lines) if lines else "—"

def _project_block_text(project: Dict[str, Any]) -> str:
    return "\n".join(f"{k}: {v}" for k, v in (project or {}).items() if v) or "—"

def _prompt_text_from_pkg(filename: str) -> str:
    return _load_pkg_text(filename)

def generate_section_via_openrouter(
    *, docx_path: str, project: Dict[str, Any], equipment: List[Dict[str, Any]],
    counts: Dict[str, Any], prompt_filename: str, cache_suffix: str,
    use_cache: bool = True, force_refresh: bool = False, pause_seconds: int = 0
) -> str:
    """
    Универсальная генерация plain-текста раздела по шаблонному промпту из пакета.
    """
    cache_path = _section_cache_path(docx_path, cache_suffix)
    if use_cache and not force_refresh and os.path.exists(cache_path):
        try:
            return open(cache_path, "r", encoding="utf-8").read()
        except Exception:
            pass

    project_block   = _project_block_text(project)
    equipment_block = _equipment_block_text(equipment)
    counts_block    = _counts_block_text(counts)

    prompt_template = _prompt_text_from_pkg(prompt_filename)
    prompt = prompt_template.format(PROJECT_BLOCK=project_block,
                                    EQUIPMENT_BLOCK=equipment_block,
                                    COUNTS_BLOCK=counts_block)

    if pause_seconds:
        time.sleep(max(0, int(pause_seconds)))

    client = OpenRouterClient(OPENROUTER_API_KEY, OPENROUTER_MODEL)
    text = client.chat([{"role": "user", "content": prompt}], temperature=0.2, json_mode=False)
    text = normalize_word_breaks(text)

    try:
        with open(cache_path, "w", encoding="utf-8") as f:
            f.write(text)
    except Exception:
        pass
    return text

def generate_methodology_via_openrouter(
    docx_path: str, project: Dict[str, Any], equipment: List[Dict[str, Any]],
    counts: Dict[str, Any], *, use_cache=True, force_refresh=False
) -> str:
    cache_path = _methodology_cache_path(docx_path)
    if use_cache and not force_refresh and os.path.exists(cache_path):
        try:
            return open(cache_path, "r", encoding="utf-8").read()
        except Exception:
            pass

    prompt = _load_pkg_text(PROMPT_METHODOLOGY).format(
        PROJECT_BLOCK=_project_block_text(project),
        EQUIPMENT_BLOCK=_equipment_block_text(equipment),
        COUNTS_BLOCK=_counts_block_text(counts),
    )
    client = OpenRouterClient(OPENROUTER_API_KEY, OPENROUTER_MODEL)
    text = normalize_word_breaks(client.chat([{"role": "user", "content": prompt}], temperature=0.2))

    try:
        with open(cache_path, "w", encoding="utf-8") as f:
            f.write(text)
    except Exception:
        pass
    return text

# Узкие обёртки
def generate_tasks_via_openrouter(docx_path, project, equipment, counts, **kw) -> str:
    return generate_section_via_openrouter(docx_path=docx_path, project=project, equipment=equipment, counts=counts,
                                           prompt_filename=PROMPT_TASKS, cache_suffix="tasks", **kw)

def generate_scope_via_openrouter(docx_path, project, equipment, counts, **kw) -> str:
    return generate_section_via_openrouter(docx_path=docx_path, project=project, equipment=equipment, counts=counts,
                                           prompt_filename=PROMPT_SCOPE, cache_suffix="scope", **kw)

def generate_procedure_via_openrouter(docx_path, project, equipment, counts, **kw) -> str:
    return generate_section_via_openrouter(docx_path=docx_path, project=project, equipment=equipment, counts=counts,
                                           prompt_filename=PROMPT_PROCEDURE, cache_suffix="procedure", **kw)

def generate_tests_via_openrouter(docx_path, project, equipment, counts, **kw) -> str:
    return generate_section_via_openrouter(docx_path=docx_path, project=project, equipment=equipment, counts=counts,
                                           prompt_filename=PROMPT_TESTS, cache_suffix="tests", **kw)


# =============================================================================
# 8) Таблицы ФЕРп (реестр/результаты)
# =============================================================================

def fill_ferp_registry_table_with_values(doc: Document, data: Dict[str, Any]) -> None:
    """Заполняет таблицу «Реестр работ по ФЕРп» по заголовкам."""
    table = find_table_by_headers(
        doc, ['№','Наименование работ','Код ФЕРп','Ед. изм.','Кол-во','Примечание'], min_match=6
    )
    if not table:
        return

    # очистить всё кроме шапки
    for row in list(table.rows)[1:]:
        table._tbl.remove(row._tr)

    counts = _compute_ferp_counts(data)
    visible = 1

    for i, code in enumerate(FERP_ORDER, start=1):
        title = FERP_TASKS.get(code, "")
        unit  = UNIT_BY_CODE.get(code, "измерение")
        note  = NOTE_BY_CODE.get(code, "")

        if i == 2 and code == "01-11-011-01":
            unit = "100 измерений"
            val  = (counts.get("01-11-011-01_all", 0) or 0) / 100.0
            qty  = _fmt_ru_decimal(val)
        elif i == 6 and code == "01-11-011-01":
            val = int(counts.get("01-11-011-01_equip", 0) or 0); qty = str(val)
        elif i == 13 and code == "01-11-011-01":
            note = "Электрическая непрерывность заземляющих соединений стоек и панелей"
            val = int(counts.get("01-11-011-01_racks", 0) or 0); qty = str(val)
        else:
            val = int(counts.get(code, 0) or 0); qty = str(val)

        if (isinstance(val, (int, float)) and val == 0) or (isinstance(val, str) and val.strip() in {"0","0,0","0.0"}):
            continue

        cells = table.add_row().cells
        cells[0].text = f"{visible}."
        cells[1].text = title
        cells[2].text = f"ФЕРп {code}"
        cells[3].text = unit
        cells[4].text = qty
        cells[5].text = note
        visible += 1

def fill_ferp_results_table_with_values(doc: Document, data: Dict[str, Any]) -> None:
    """Заполняет таблицу результатов ПНР (с датами/подписями — пока пустые поля)."""
    table = find_table_by_headers(
        doc,
        ['№','Наименование работ','Код ФЕРп','Ед. изм.','Кол-во','Дата',
         'Результат проверки (OK / DEFECT)','Ф.И.О. исполнителя','Подпись'],
        min_match=9
    )
    if not table:
        return

    for row in list(table.rows)[1:]:
        table._tbl.remove(row._tr)

    counts = _compute_ferp_counts(data)

    def _row_title(i: int, code: str) -> str:
        if i == 2 and code == "01-11-011-01":
            return "Проверка электрической непрерывности между заземлителями и заземлёнными элементами оборудования"
        if i == 5 and code == "02-01-002-01":
            return "Проверка целостности и параметров линий связи (электрических или оптических)"
        return FERP_TASKS.get(code, "")

    visible = 1
    for i, code in enumerate(FERP_ORDER, start=1):
        title = _row_title(i, code)
        unit  = UNIT_BY_CODE.get(code, "измерение")

        if i == 2 and code == "01-11-011-01":
            unit = "100 измерений"
            value = (counts.get("01-11-011-01_all", 0) or 0) / 100.0
            qty = _fmt_ru_decimal(value)
        elif i == 6 and code == "01-11-011-01":
            value = counts.get("01-11-011-01_equip", 0) or 0; qty = str(value)
        elif i == 13 and code == "01-11-011-01":
            value = counts.get("01-11-011-01_racks", 0) or 0; qty = str(value)
        else:
            value = counts.get(code, 0) or 0; qty = str(value)

        if not value or (isinstance(value, str) and value.strip() in {"0","0,0","0.0"}):
            continue

        cells = table.add_row().cells
        cells[0].text = f"{visible}."
        cells[1].text = title
        cells[2].text = f"ФЕРп {code}"
        cells[3].text = unit
        cells[4].text = qty
        cells[5].text = ""  # дата
        cells[6].text = ""  # результат
        cells[7].text = ""  # ФИО
        cells[8].text = ""  # подпись
        visible += 1


# =============================================================================
# 9) Сборка DOCX по шаблону
# =============================================================================

def build_docx_with_tpl(template_path: str, output_path: str, data: Dict[str, Any]) -> str:
    """Рендерит итоговый DOCX по шаблону и данным."""
    doc_tpl = DocxTemplate(template_path)

    # исходные данные
    project  = data.get("project", {}) or {}
    eq_list  = data.get("equipment", []) or []
    pnr      = data.get("pnr", {}) or {}
    counts   = _compute_ferp_counts(data)
    qty_100  = _fmt_ru_decimal((counts.get("01-11-011-01_all", 0) or 0) / 100.0)

    tasks    = [strip_num_prefix(s) for s in (pnr.get("main_tasks_5") or [])]
    checklist= pnr.get("checklist_10_15") or []

    # базовый контекст шаблона
    context = {
        "project": project,
        "tasks": [f"{i+1}) {t}" for i, t in enumerate(tasks)],
        "tasks_block": "\n".join(f"{i+1}) {t}" for i, t in enumerate(tasks)),
        "checklist": checklist,
        "pnr_block": "\n".join(checklist),
        "methodology_block": pnr.get("methodology_section", "") or "",
        "total": int(data.get("summaries", {}).get("total_switch_ports", 0)),
        "ferp_qty": {
            "01-11-010-01": counts.get("01-11-010-01") or "0. Не применяется",
            "01-11-028-01": counts.get("01-11-028-01") or "0. Не применяется",
            "01-06-022-01": counts.get("01-06-022-01") or "0. Не применяется",
            "02-01-002-01": counts.get("02-01-002-01") or "0. Не применяется",
            "02-01-003-01": counts.get("02-01-003-01") or "0. Не применяется",
            "02-01-004-01": counts.get("02-01-004-01") or "0. Не применяется",
            "01-11-012-01": counts.get("01-11-012-01") or "0. Не применяется",
            "01-11-026-01": counts.get("01-11-026-01") or "0. Не применяется",
            "01-11-034-01": counts.get("01-11-034-01") or "0. Не применяется",
            "01-11-011-02": counts.get("01-11-011-02") or "0. Не применяется",
            "01-11-029-01": counts.get("01-11-029-01") or "0. Не применяется",
            "01-11-036-01": counts.get("01-11-036-01") or "0. Не применяется",
            "01-11-011-01_all":   counts.get("01-11-011-01_all", 0),
            "01-11-011-01_equip": counts.get("01-11-011-01_equip", 0),
            "01-11-011-01_racks": counts.get("01-11-011-01_racks", 0),
        },
        "ferp_qty_100_measure": qty_100,

        # LLM-разделы (подставим позже)
        "tasks_block_full": "",
        "scope_block_full": "",
        "procedure_block_full": "",
        "tests_block_full": "",
    }

    # базовый путь исходника (для кэша разделов)
    src_docx = output_path.replace("_pnr.docx", ".docx")

    # генерим разделы через LLM (если пусто в context)
    if not context["tasks_block_full"]:
        context["tasks_block_full"] = generate_tasks_via_openrouter(src_docx, project, eq_list, counts)
    if not context["scope_block_full"]:
        context["scope_block_full"] = generate_scope_via_openrouter(src_docx, project, eq_list, counts)
    if not context["procedure_block_full"]:
        context["procedure_block_full"] = generate_procedure_via_openrouter(src_docx, project, eq_list, counts)
    if not context["tests_block_full"]:
        context["tests_block_full"] = generate_tests_via_openrouter(src_docx, project, eq_list, counts)

    # методология (если пустая)
    if not context["methodology_block"]:
        context["methodology_block"] = generate_methodology_via_openrouter(
            docx_path=src_docx, project=project, equipment=eq_list, counts=counts, use_cache=True
        )

    # рендер контекста в шаблон
    doc_tpl.render(context)
    doc_tpl.save(output_path)

    # «живая» вставка абзацев/списков по плейсхолдерам
    doc = Document(output_path)
    ensure_paragraph_style(doc, "PNR_Body", font_name="Cambria", font_size_pt=12, left_indent_pt=18)

    for key in ("tasks_block_full", "scope_block_full", "procedure_block_full", "tests_block_full", "methodology_block"):
        context[key] = normalize_word_breaks(context.get(key, ""))

    insert_text_with_style(doc, "{{tasks_block_full}}",     context["tasks_block_full"],     "PNR_Body")
    insert_text_with_style(doc, "{{scope_block_full}}",     context["scope_block_full"],     "PNR_Body")
    insert_text_with_style(doc, "{{procedure_block_full}}", context["procedure_block_full"], "PNR_Body")
    insert_text_with_style(doc, "{{tests_block_full}}",     context["tests_block_full"],     "PNR_Body")
    insert_text_with_style(doc, "{{methodology_block}}",    context["methodology_block"],    "PNR_Body")

    # Таблицы
    _fill_equipment_and_channels_tables(doc, eq_list)
    fill_ferp_registry_table_with_values(doc, data)
    fill_ferp_results_table_with_values(doc, data)

    # Помечаем поля (TOC/REF/DOCPROPERTY) «грязными», чтобы Word предложил обновить
    try:
        for field in doc.element.findall(".//w:fldSimple", doc.element.nsmap):
            instr = field.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}instr")
            if instr and any(k in instr for k in ("REF", "DOCPROPERTY", "SEQ", "TOC")):
                field.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}dirty", "true")
    except Exception:
        pass

    doc.save(output_path)
    return output_path

def _fill_equipment_and_channels_tables(doc: Document, eq_list: List[Dict[str, Any]]) -> None:
    """Заполняет таблички «Оборудование» и «Каналы», если они есть в шаблоне."""
    # Оборудование
    equipment_table = find_table_by_headers(doc, ['№', 'Наименование', 'Тип/марка', 'Работы'], min_match=3)
    if equipment_table:
        for row in equipment_table.rows[1:2]:
            if all(not cell.text.strip() for cell in row.cells):
                equipment_table._tbl.remove(row._tr)
        for idx, item in enumerate(eq_list, start=1):
            cells = equipment_table.add_row().cells
            cells[0].text = str(idx)
            cells[1].text = item.get('name', '')
            cells[2].text = item.get('model_or_code', '—')
            cells[3].text = item.get('work', '')

    # Каналы
    channels_table = find_table_by_headers(doc, ['Количество каналов на ед.'], min_match=1)
    if channels_table:
        for row in channels_table.rows[1:2]:
            if all(not cell.text.strip() for cell in row.cells):
                channels_table._tbl.remove(row._tr)

        grand_total = 0
        row_idx = 1
        for it in eq_list:
            et = (it.get('equipment_type') or '').lower()
            if et not in {'switch', 'router'}:
                continue
            qty = int(it.get('quantity', 1))
            pc  = int(it.get('port_count') or 0)
            total = qty * (pc + 1)  # + mgmt
            grand_total += total

            cells = channels_table.add_row().cells
            cells[0].text = str(row_idx)
            cells[1].text = it.get('name', '—')
            cells[2].text = str(qty)
            cells[3].text = str(pc + 1)
            cells[4].text = str(total)
            row_idx += 1

        extra = sum(int(x.get('quantity') or 0) for x in eq_list if (x.get('equipment_type') or '').lower() in {'voip_phone', 'ups'})
        if extra:
            cells = channels_table.add_row().cells
            cells[0].text = str(row_idx)
            cells[1].text = "Дополнительные порты (телефоны, ИБП)"
            cells[2].text = "—"
            cells[3].text = "—"
            cells[4].text = str(extra)
            grand_total += extra

        cells = channels_table.add_row().cells
        try:
            cells[0].merge(cells[3]).text = "Итого"
        except Exception:
            cells[0].text = "Итого"
        cells[4].text = str(grand_total)


# =============================================================================
# 10) Конвейер
# =============================================================================

def normalize_equipment(equipment: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """Мерджим дубли по (equipment_type, vendor, model_or_code/name) + суммируем quantity."""
    def key(it: Dict[str, Any]) -> str:
        return "|".join([
            (it.get("equipment_type") or "").lower(),
            (it.get("vendor") or "").lower(),
            ((it.get("model_or_code") or it.get("name") or "")).lower()
        ])
    merged: Dict[str, Dict[str, Any]] = {}
    for it in equipment or []:
        q = it.get("quantity", 1)
        try:
            q = int(q)
        except Exception:
            q = 1
        it["quantity"] = q
        k = key(it)
        if k in merged:
            merged[k]["quantity"] += q
            try:
                merged[k]["port_count"] = max(int(merged[k].get("port_count") or 0), int(it.get("port_count") or 0))
            except Exception:
                pass
        else:
            merged[k] = dict(it)
    return list(merged.values())

def call_gemini_structured(docx_text: str, docx_path: str, *, use_cache=True, force_refresh=False) -> Dict[str, Any]:
    """
    Совместимая «обёртка» (старое имя функции). Берёт из кэша или вызывает OpenRouter.
    """
    cache_path = gemini_cache_path(docx_path)
    if use_cache and not force_refresh and os.path.exists(cache_path):
        try:
            with open(cache_path, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            pass

    data = call_structured_extractor(docx_text, docx_path)

    try:
        with open(cache_path, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
    except Exception:
        pass
    return data

def process_one_docx_file(docx_path: str, template_path: str) -> Optional[str]:
    """Полный цикл обработки одного .docx → новый _pnr.docx."""
    if not os.path.exists(docx_path):
        print(f"[ERROR] Файл не найден: {docx_path}"); return None
    if not os.path.exists(template_path):
        print(f"[ERROR] Шаблон не найден: {template_path}"); return None

    print(f"[READ] {docx_path}")
    try:
        docx_text = read_docx_text(docx_path)
    except Exception as e:
        print(f"[WARN] Не удалось прочитать DOCX: {e}"); return None
    if not docx_text.strip():
        print("[WARN] Пустой текст в документе."); return None

    print("[LLM] Извлечение JSON по схеме через OpenRouter…")
    try:
        data = call_gemini_structured(docx_text, docx_path)
        data["equipment"] = normalize_equipment(data.get("equipment", []))
        data.setdefault("summaries", {"total_switch_ports": 0, "notes": ""})
        data.setdefault("pnr", {"main_tasks_5": [], "checklist_10_15": [], "methodology_section": ""})
    except Exception as e:
        print(f"[ERROR] Structured extract: {e}"); return None

    out_path = os.path.splitext(docx_path)[0] + "_pnr.docx"
    print(f"[DOCX] Рендер → {out_path}")
    try:
        build_docx_with_tpl(template_path, out_path, data)
    except Exception as e:
        print(f"[ERROR] DOCX сборка: {e}"); return None

    return out_path

def main():
    path = DOCX_PATH_DEFAULT
    template_path = DOCX_TEMPLATE_PATH_DEFAULT

    if os.path.isfile(path):
        res = process_one_docx_file(path, template_path)
        print(json.dumps({"processed": 1 if res else 0, "files": [res] if res else []}, ensure_ascii=False, indent=2))
        return

    if os.path.isdir(path):
        docx_files = sorted(
            f for f in (os.path.join(path, n) for n in os.listdir(path))
            if os.path.isfile(f)
            and f.lower().endswith(".docx")
            and not os.path.basename(f).startswith("~$")
            and not f.lower().endswith("_pnr.docx")
        )
        built: List[str] = []
        for p in docx_files:
            print("=" * 80)
            print(f"Файл: {os.path.basename(p)}")
            r = process_one_docx_file(p, template_path)
            if r:
                built.append(r)
        print("=" * 80)
        print(json.dumps({"processed": len(built), "files": built}, ensure_ascii=False, indent=2))
        return

    print(f"[ERROR] Путь не найден: {path}")

if __name__ == "__main__":
    main()
